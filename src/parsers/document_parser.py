"""
Document parser using the unstructured library to extract structured elements.
Handles multiple file formats and converts them into DocumentElement objects.
"""

import io
import tempfile
import os
from typing import List, Dict, Any, Optional, Union
from pathlib import Path

from unstructured.partition.auto import partition
from unstructured.documents.elements import Element

from src.models.resume_elements import (
    DocumentElement,
    ElementType,
    ParsedDocument,
    ResumeSection
)


class DocumentParser:
    """Handles document parsing using the unstructured library."""
    
    def __init__(self):
        """Initialize the document parser."""
        self.supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.html', '.htm'}
    
    def parse_document(
        self,
        file_content: bytes,
        filename: str,
        mime_type: Optional[str] = None
    ) -> ParsedDocument:
        """
        Parse a document and extract structured elements.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            mime_type: MIME type if available
            
        Returns:
            ParsedDocument with extracted elements
        """
        file_extension = Path(filename).suffix.lower()
        file_type = self._get_file_type(file_extension)
        
        try:
            # Create temporary file for unstructured processing
            with tempfile.NamedTemporaryFile(
                suffix=file_extension,
                delete=False
            ) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Set environment variables for OpenGL fallback
                import os
                os.environ.setdefault('LIBGL_ALWAYS_SOFTWARE', '1')
                os.environ.setdefault('MESA_GL_VERSION_OVERRIDE', '3.3')
                
                elements = self._partition_document(temp_file_path)
                
                # Check if the parsing extracted meaningful content
                meaningful_content = False
                if elements:
                    total_text = ' '.join([str(elem) for elem in elements])
                    # Check if we have substantial text content (not just empty or whitespace)
                    meaningful_text = total_text.strip()
                    # Check for meaningful content (not just dots, empty strings, or very short content)
                    if (len(meaningful_text) > 50 and 
                        not meaningful_text.replace('.', '').replace(' ', '').replace('\n', '') == '' and
                        not all(char in ' .\n\t' for char in meaningful_text)):
                        meaningful_content = True
                
                if not meaningful_content:
                    # If unstructured didn't extract meaningful content, use fallback
                    return self._parse_with_fallback(file_content, filename, file_extension, file_type)
                
                # Convert unstructured elements to our DocumentElement format
                document_elements = self._convert_elements(elements)
                
                # Create parsed document
                parsed_doc = ParsedDocument(
                    filename=filename,
                    file_extension=file_extension,
                    file_type=file_type,
                    total_elements=len(document_elements),
                    grouped_sections=[],  # Will be populated by ContentProcessor
                    parsing_warnings=[]
                )
                
                # Store the elements for processing
                parsed_doc._raw_elements = document_elements
                
                return parsed_doc
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            error_msg = str(e)
            
            # Handle specific OpenGL errors
            if "libGL.so.1" in error_msg or "OpenGL" in error_msg:
                # Try fallback approach for Railway deployment
                try:
                    return self._parse_with_fallback(file_content, filename, file_extension, file_type)
                except Exception as fallback_error:
                    return ParsedDocument(
                        filename=filename,
                        file_extension=file_extension,
                        file_type=file_type,
                        total_elements=0,
                        grouped_sections=[],
                        parsing_warnings=[
                            f"Parsing failed due to OpenGL library issues: {error_msg}",
                            f"Fallback parsing also failed: {str(fallback_error)}"
                        ]
                    )
            
            # For other errors, try fallback parsing
            try:
                return self._parse_with_fallback(file_content, filename, file_extension, file_type)
            except Exception as fallback_error:
                return ParsedDocument(
                    filename=filename,
                    file_extension=file_extension,
                    file_type=file_type,
                    total_elements=0,
                    grouped_sections=[],
                    parsing_warnings=[f"Parsing failed: {error_msg}"]
                )
    
    def _convert_elements(self, unstructured_elements: List[Element]) -> List[DocumentElement]:
        """
        Convert unstructured library elements to our DocumentElement format.
        
        Args:
            unstructured_elements: Elements from unstructured library
            
        Returns:
            List of DocumentElement objects
        """
        document_elements = []
        
        for element in unstructured_elements:
            try:
                # Map unstructured element types to our ElementType enum
                element_type = self._map_element_type(element.category)
                
                # Extract metadata
                metadata = {}
                if hasattr(element, 'metadata') and element.metadata:
                    try:
                        metadata = element.metadata.to_dict() if hasattr(element.metadata, 'to_dict') else {}
                    except:
                        metadata = {}
                
                # Extract page number if available
                page_number = None
                if 'page_number' in metadata:
                    page_number = metadata.get('page_number')
                
                # Extract coordinates if available - handle unstructured coordinate objects
                coordinates = None
                if 'coordinates' in metadata:
                    try:
                        coord_data = metadata.get('coordinates')
                        # Only include coordinates if they're simple numeric values
                        if isinstance(coord_data, dict) and all(isinstance(v, (int, float)) for v in coord_data.values()):
                            coordinates = coord_data
                    except:
                        coordinates = None
                
                # Get text content safely
                try:
                    text_content = str(element)
                except:
                    text_content = ""
                
                # Skip empty elements
                if not text_content.strip():
                    continue
                
                # Create DocumentElement
                doc_element = DocumentElement(
                    element_type=element_type,
                    text=text_content,
                    metadata=metadata,
                    page_number=page_number,
                    coordinates=coordinates
                )
                
                document_elements.append(doc_element)
                
            except Exception as e:
                # Skip problematic elements but log the issue
                continue
        
        return document_elements
    
    def _map_element_type(self, unstructured_category: str) -> ElementType:
        """
        Map unstructured library element categories to our ElementType enum.
        
        Args:
            unstructured_category: Category from unstructured library
            
        Returns:
            Corresponding ElementType
        """
        category_mapping = {
            'Title': ElementType.TITLE,
            'NarrativeText': ElementType.NARRATIVE_TEXT,
            'ListItem': ElementType.LIST_ITEM,
            'Table': ElementType.TABLE,
            'Header': ElementType.HEADER,
            'Footer': ElementType.FOOTER,
            'EmailAddress': ElementType.EMAIL_ADDRESS,
            'Address': ElementType.ADDRESS,
            'PhoneNumber': ElementType.PHONE_NUMBER,
            'UncategorizedText': ElementType.NARRATIVE_TEXT,  # Fallback
            'Text': ElementType.NARRATIVE_TEXT,  # Fallback
        }
        
        return category_mapping.get(unstructured_category, ElementType.NARRATIVE_TEXT)
    
    def _get_file_type(self, extension: str) -> str:
        """
        Get human-readable file type from extension.
        
        Args:
            extension: File extension (e.g., '.pdf')
            
        Returns:
            Human-readable file type
        """
        type_mapping = {
            '.pdf': 'PDF Document',
            '.docx': 'Word Document (DOCX)',
            '.doc': 'Word Document (DOC)',
            '.txt': 'Text Document',
            '.html': 'HTML Document',
            '.htm': 'HTML Document'
        }
        
        return type_mapping.get(extension, 'Unknown Document Type')
    
    def get_document_stats(self, parsed_doc: ParsedDocument) -> Dict[str, Any]:
        """
        Get statistics about the parsed document.
        
        Args:
            parsed_doc: Parsed document
            
        Returns:
            Dictionary with document statistics
        """
        element_counts = {}
        total_text_length = 0
        
        for section in parsed_doc.grouped_sections:
            for element in section.elements:
                element_type = element.element_type.value
                element_counts[element_type] = element_counts.get(element_type, 0) + 1
                total_text_length += len(element.text)
        
        return {
            'total_elements': parsed_doc.total_elements,
            'element_type_counts': element_counts,
            'total_text_length': total_text_length,
            'has_warnings': len(parsed_doc.parsing_warnings) > 0,
            'warning_count': len(parsed_doc.parsing_warnings)
        }

    def _partition_document(self, temp_file_path: str):
        """
        Partition document using unstructured library.
        This method can be overridden for testing purposes.
        
        Args:
            temp_file_path: Path to temporary file
            
        Returns:
            List of unstructured elements
        """
        from unstructured.partition.auto import partition
        
        return partition(
            filename=temp_file_path,
            strategy="auto",
            include_page_breaks=True,
            infer_table_structure=True,
            chunking_strategy=None
        )

    def _parse_with_fallback(self, file_content: bytes, filename: str, file_extension: str, file_type: str) -> ParsedDocument:
        """
        Fallback parsing method when OpenGL libraries are not available.
        Uses alternative approaches for document parsing.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            file_extension: File extension
            file_type: Human-readable file type
            
        Returns:
            ParsedDocument with extracted elements
        """
        try:
            # For text files, use simple text extraction
            if file_extension in ['.txt']:
                text_content = file_content.decode('utf-8', errors='ignore')
                
                # Split text into lines and create structured elements
                lines = text_content.split('\n')
                elements = []
                
                for i, line in enumerate(lines):
                    line = line.strip()
                    if line:
                        # Determine element type based on content
                        if any(keyword in line.lower() for keyword in ['summary', 'objective', 'profile', 'about']):
                            element_type = ElementType.TITLE
                        elif any(keyword in line.lower() for keyword in ['experience', 'work', 'employment', 'career']):
                            element_type = ElementType.TITLE
                        elif any(keyword in line.lower() for keyword in ['education', 'academic', 'degree', 'university']):
                            element_type = ElementType.TITLE
                        elif any(keyword in line.lower() for keyword in ['skills', 'competencies', 'expertise', 'technologies']):
                            element_type = ElementType.TITLE
                        elif any(keyword in line.lower() for keyword in ['certifications', 'certificates', 'licenses']):
                            element_type = ElementType.TITLE
                        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                            element_type = ElementType.LIST_ITEM
                        elif '@' in line and '.' in line:  # Email address
                            element_type = ElementType.EMAIL_ADDRESS
                        elif any(char.isdigit() for char in line) and any(char in '()-' for char in line):  # Phone number
                            element_type = ElementType.PHONE_NUMBER
                        else:
                            element_type = ElementType.NARRATIVE_TEXT
                        
                        elements.append(DocumentElement(
                            element_type=element_type,
                            text=line,
                            metadata={'line_number': i + 1},
                            page_number=1
                        ))
                
                parsed_doc = ParsedDocument(
                    filename=filename,
                    file_extension=file_extension,
                    file_type=file_type,
                    total_elements=len(elements),
                    grouped_sections=[],
                    parsing_warnings=["Document parsed successfully using text extraction"]
                )
                parsed_doc._raw_elements = elements
                return parsed_doc
            
            # For PDF files, try using pdfplumber as fallback
            elif file_extension == '.pdf':
                try:
                    import pdfplumber
                    import io
                    
                    with io.BytesIO(file_content) as pdf_stream:
                        with pdfplumber.open(pdf_stream) as pdf:
                            elements = []
                            for page_num, page in enumerate(pdf.pages, 1):
                                text = page.extract_text()
                                if text:
                                    # Split text into lines and create structured elements
                                    lines = text.split('\n')
                                    for i, line in enumerate(lines):
                                        line = line.strip()
                                        if line:
                                            # Determine element type based on content
                                            if any(keyword in line.lower() for keyword in ['summary', 'objective', 'profile', 'about']):
                                                element_type = ElementType.TITLE
                                            elif any(keyword in line.lower() for keyword in ['experience', 'work', 'employment', 'career']):
                                                element_type = ElementType.TITLE
                                            elif any(keyword in line.lower() for keyword in ['education', 'academic', 'degree', 'university']):
                                                element_type = ElementType.TITLE
                                            elif any(keyword in line.lower() for keyword in ['skills', 'competencies', 'expertise', 'technologies']):
                                                element_type = ElementType.TITLE
                                            elif any(keyword in line.lower() for keyword in ['certifications', 'certificates', 'licenses']):
                                                element_type = ElementType.TITLE
                                            elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
                                                element_type = ElementType.LIST_ITEM
                                            elif '@' in line and '.' in line:  # Email address
                                                element_type = ElementType.EMAIL_ADDRESS
                                            elif any(char.isdigit() for char in line) and any(char in '()-' for char in line):  # Phone number
                                                element_type = ElementType.PHONE_NUMBER
                                            else:
                                                element_type = ElementType.NARRATIVE_TEXT
                                            
                                            elements.append(DocumentElement(
                                                element_type=element_type,
                                                text=line,
                                                metadata={'page_number': page_num, 'line_number': i + 1},
                                                page_number=page_num
                                            ))
                    
                    parsed_doc = ParsedDocument(
                        filename=filename,
                        file_extension=file_extension,
                        file_type=file_type,
                        total_elements=len(elements),
                        grouped_sections=[],
                        parsing_warnings=["Document parsed successfully using PDF text extraction"]
                    )
                    parsed_doc._raw_elements = elements
                    return parsed_doc
                except ImportError:
                    # If pdfplumber is not available, try basic text extraction
                    pass
            
            # For DOCX files, try using python-docx as fallback
            elif file_extension == '.docx':
                try:
                    from docx import Document
                    import io
                    
                    with io.BytesIO(file_content) as docx_stream:
                        doc = Document(docx_stream)
                        elements = []
                        
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                elements.append(DocumentElement(
                                    element_type=ElementType.NARRATIVE_TEXT,
                                    text=paragraph.text,
                                    metadata={},
                                    page_number=1
                                ))
                    
                    parsed_doc = ParsedDocument(
                        filename=filename,
                        file_extension=file_extension,
                        file_type=file_type,
                        total_elements=len(elements),
                        grouped_sections=[],
                        parsing_warnings=["Document parsed successfully using Word document extraction"]
                    )
                    parsed_doc._raw_elements = elements
                    return parsed_doc
                except ImportError:
                    # If python-docx is not available, try basic text extraction
                    pass
            
            # Ultimate fallback: try to extract any readable text
            try:
                text_content = file_content.decode('utf-8', errors='ignore')
                if not text_content.strip():
                    # Try other encodings
                    for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                        try:
                            text_content = file_content.decode(encoding, errors='ignore')
                            if text_content.strip():
                                break
                        except:
                            continue
                
                if text_content.strip():
                    elements = [DocumentElement(
                        element_type=ElementType.NARRATIVE_TEXT,
                        text=text_content,
                        metadata={},
                        page_number=1
                    )]
                    
                    parsed_doc = ParsedDocument(
                        filename=filename,
                        file_extension=file_extension,
                        file_type=file_type,
                        total_elements=len(elements),
                        grouped_sections=[],
                        parsing_warnings=["Used basic text extraction fallback due to OpenGL library issues"]
                    )
                    parsed_doc._raw_elements = elements
                    return parsed_doc
            except:
                pass
            
            # If all fallbacks fail, return empty document with error
            raise Exception("All fallback parsing methods failed")
            
        except Exception as e:
            raise Exception(f"Fallback parsing failed: {str(e)}")
